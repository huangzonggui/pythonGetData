# -*- coding: utf-8 -*-

#author:王得伟Dewitt
#email:wangdewei1996@163.com
#date:2017.09.05

from urllib import request
import os  
import re
import sys
from bs4 import BeautifulSoup
import pdfkit
from docx import Document  
from docx.shared import Inches 
from docx.shared import Pt
import time  # 引入time模块

host = "sports.qq.com"

localtime = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())#日期
dateStr = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())
# 格式化成2016-03-20 11:45:39形式
#localtime = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()) 
doc=Document()  #新建文档
##word title  
doc.add_paragraph(dateStr)
text = doc.add_paragraph('NBA新闻')
#doc.add_heading('腾讯财经',0)

def wirteToDoc():
    '''
    向后台发起请求，把需要的内容写入docx文档中
    '''
    headers = {#为什么将这个头放在while外面不行？？？第二个以后的请求都是Access defined
        "Referer" : "http://" + host,   #需要添加Referer头部，否则请求失败
        "Accept-Encoding": ""
    }
    req = request.Request(
        "http://sports.qq.com/basket/nba/morenews.htm",
                          headers = headers
    )
    response = request.urlopen(req)
    jsonData = response.read().decode("gbk")
    jsonData = jsonData.replace('\\', '')
    #print(jsonData)
    ret = re.findall(r"<span.+?</a>", jsonData)
    #print(ret)#ret是新闻链接URL list

    for one in ret:
        bsObj = BeautifulSoup(one, "html.parser")
        title = bsObj.find(
            "a",
            {"class": "fs14"}
        )
        print(title)

        timeStr = bsObj.find(
            "span",
            {"class": "flr"}    
        )
        print(timeStr)

        if title != None:            
            text.add_run('\n'+str(title.string))
            run = text.add_run(str(timeStr.string))
            run.font.size = Pt(9)
            run.italic = True#斜体

if __name__ == "__main__":
    wirteToDoc()#把标题时间抽取出来放在doc中
    #word  
    saveFile=os.getcwd()+"\\NBA新闻"+localtime+".docx"  
    doc.save(saveFile)#根据saveFile的路径和文件名保存文件
