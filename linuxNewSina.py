from urllib import request
import os  
import re
import sys
from bs4 import BeautifulSoup
import pdfkit
from docx import Document  
from docx.shared import Inches 
from docx.shared import Pt
import time  
import json
import ssl

ssl._create_default_https_context = ssl._create_unverified_context

#host = "roll.news.sina.com.cn"
host = "feed.mix.sina.com.cn"
pageNum = 1

localtime = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())#日期
dateStr = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime()) 
doc=Document()  #新建文档x
#doc.add_paragraph(dateStr)
text = doc.add_paragraph('')
def getSinaNews():
    global pageNum
    headers = {
        "Accept":"*/*",
        "Accept-Encoding":"gzip, deflate",
        "Accept-Language":"zh-CN,zh;q=0.8",
        "Connection":"keep-alive",
        "Host": host,
        "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/60.0.3112.113 Safari/537.36",
        "Referer" : "http://news.sina.com.cn/roll/",   #需要添加Referer头部，否则请求失败
        "Accept-Encoding": ""
    }
    
    while  pageNum <= 350 :
        print (pageNum)
        req = request.Request(
        "https://"+ host +"/api/roll/get?pageid=153&lid=2509&k=&num=50&page="+str(pageNum)+"&r=0.7447433997811357",
        )
        
        response = request.urlopen(req)

        jsonData = response.read(500000).decode("unicode_escape")
        #print (jsonData)
        ret = re.findall(r",\"title\":\".*?,", jsonData)#?非贪婪匹配 
        times = re.findall(r"\"ctime\":\".*?,", jsonData)
        #print (ret)
        #print (times)


        i = 0
        for item in ret: 
            #print (item)

            run = text.add_run('\n'+ item[10:-2] + '  ')

            #print (item[10:-2])
            run.font.bold = True#加粗
            run.font.size = Pt(9)

            #print (times[i])
            if i <= len(times) - 1:
                timearray = time.localtime(int(times[i][9:-2]))
                otherstyletime = time.strftime("%m-%d %H:%M", timearray)
                run = text.add_run(otherstyletime)
                run.font.size = Pt(8)
                #print (otherstyletime)
            i=i+1
       
        pageNum = pageNum + 1
        #time.sleep(1)
    return ret


if __name__ == "__main__":
    try:
        getSinaNews()
    finally: 
        saveFile="sina"+localtime+".docx"  
        doc.save(saveFile)#根据saveFile的路径和文件名保存文件

    
