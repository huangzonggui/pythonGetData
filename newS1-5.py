from urllib import request
import os  
import re
import sys
from bs4 import BeautifulSoup
import pdfkit
from docx import Document  
from docx.shared import Inches 
from docx.shared import Pt
import time  # 引入time模块

host = "roll.news.sina.com.cn"

localtime = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())#日期
dateStr = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())
doc=Document()  #新建文档x
doc.add_paragraph(dateStr)
text = doc.add_paragraph('新浪')

def getNewsHrefs():
    '''
    通过ajax请求获取首页新闻链接列表
    :return: 首页前50条新闻链接URL list
    '''
    headers = {
        "Referer" : "http://" + host,   #需要添加Referer头部，否则请求失败
        "Accept-Encoding": ""
    }
    req = request.Request(
        "http://"+ host +"/interface/rollnews_ch_out_interface.php?col=89&spec=&type=&date=&ch=01&k=&offset_page=0&offset_num=0&num=60&asc=&page=1&last_time=1516540224&r=0.5169514536772964",
		                      headers = headers
    )
    response = request.urlopen(req)
    jsonData = response.read(50000).decode("gbk")
    # print (jsonData)
    jsonData = jsonData.replace('\\', '')

    ret = re.findall(r"\"(http://.+?)\"", jsonData)#?非贪婪匹配

    #第2页到下面的页数
    i=2
    while i <= 2:
        req = request.Request(
        "http://" + host + "/interface/rollnews_ch_out_interface.php?col=43&spec=&type=&ch=03&k=&offset_page=0&offset_num=0&num=60&asc=&page="+str(i)+"&r=0.7229786096128036",
                              headers = headers
        )
        response = request.urlopen(req)
        jsonData = response.read(50000).decode("gbk")
        jsonData = jsonData.replace('\\', '')
        #print (jsonData)
        ret = ret + re.findall(r"\"(http://.+?)\"", jsonData)#?非贪婪匹配
        i = i + 1
    return ret


def getNews(newsHrefs):
    '''
    将新闻链接列表中所有新闻内容提取并拼凑成一个新的页面
    :param newsHrefs:新闻链接URL list
    :return: 包含所有新闻内容的html string
    '''
    ret = '''
            <!DOCTYPE html><html lang="zh-CN">
            <head>
            <meta content="text/html; charset=utf-8" http-equiv="Content-Type">
            <meta charset="utf-8">            
            </head>
            <body>
            '''

    headers = {
        "Accept-Encoding": ""
    }
    j=0
    k=0
    for href in newsHrefs:
        j = j+1
        req = request.Request(
            href,
            headers=headers
        )
        response = request.urlopen(req)

        #TODO:存在某些页面无法解码，原因不明，暂时忽略
        try:
            html = response.read().decode("utf8", "ignore")
        except (ValueError, IndexError) as e:
            pass
        bsObj = BeautifulSoup(html, "html.parser")

        page = bsObj.find(#包裹时间跟标题的div
            "div",
            {"class": "main-content"}#网站升级的时候导致出问题
        )
        if page:
            k = k+1
            title = page.find(
                "h1",
                {"class": "main-title"}
            )
            timeStr = page.find(
			    "span",
                {"class": "date"}#美股
            )
            print (title)
            #ret += (str(title) + str(timeStr))
            ret += ('str')
            text.add_run('\n'+str(title.string) + '  ')
            run = text.add_run(str(timeStr.string)[6:7]+'-'+str(timeStr.string)[8:10]+'  '+str(timeStr.string)[11:17])
            run.font.size = Pt(9)
            run.italic = True#斜体

        #print('共'+j+'条;'+'拿到了'+k+'条')
        if href == newsHrefs[5]:
            break
    ret += '''
    </body>
    </html>
    '''
    #print (ret)
    return ret

if __name__ == "__main__":
    newsHrefs = getNewsHrefs()
    #print (newsHrefs)
    news = getNews(newsHrefs)

    saveFile=os.getcwd()+"\\新浪"+localtime+".docx"  
    doc.save(saveFile)#根据saveFile的路径和文件名保存文件
