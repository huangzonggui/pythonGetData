from urllib import request
import os  
import re
import sys
from bs4 import BeautifulSoup
import pdfkit
from docx import Document  
from docx.shared import Inches 
from docx.shared import Pt
import time  # 引入time模块
import json

host = "roll.news.sina.com.cn"
pageNum = 1

localtime = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())#日期
dateStr = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())
doc=Document()  #新建文档x
#doc.add_paragraph(dateStr)
text = doc.add_paragraph('新浪')

def getSinaNews():
    global pageNum
    headers = {
        "Accept":"*/*",
        "Accept-Encoding":"gzip, deflate",
        "Accept-Language":"zh-CN,zh;q=0.8",
        "Connection":"keep-alive",
        "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/60.0.3112.113 Safari/537.36",
        "Referer" : "http://" + host,   #需要添加Referer头部，否则请求失败
        "Accept-Encoding": ""
    }
    
    while  pageNum <= 160:
        print (pageNum)
        req = request.Request(
        "http://" + host + "/interface/rollnews_ch_out_interface.php?col=89&spec=&type=&ch=01&k=&offset_page=0&offset_num=0&num=60&asc=&page="+str(pageNum)+"&r=0.6575086962964121",
                              headers = headers
        )
        response = request.urlopen(req)
        jsonData = response.read(50000).decode("gbk")
        ret = re.findall(r",title : \".*?,", jsonData)#?非贪婪匹配
        times = re.findall(r",time : .*?}", jsonData)

        i = 0
        for item in ret: 
            #print (item)

            run = text.add_run('\n'+ item[10:-2] + '  ')
            #print (item[10:-2])
            run.font.bold = True#加粗
            run.font.size = Pt(9)

            timearray = time.localtime(int(times[i][8:-1]))
            otherstyletime = time.strftime("%m-%d %H:%M", timearray)
            run = text.add_run(otherstyletime)
            run.font.size = Pt(8)
            #print (otherstyletime)
            i=i+1
       
        pageNum = pageNum + 1
        time.sleep(1)
    return ret


if __name__ == "__main__":
    try:
        getSinaNews()
    except:
        saveFile=os.getcwd()+"\\新浪"+localtime+".docx"  
        doc.save(saveFile)#根据saveFile的路径和文件名保存文件
        print('error1..................................................')
        try:
            time.sleep(10)
            getSinaNews()
        except:
            saveFile=os.getcwd()+"\\新浪"+localtime+".docx"  
            doc.save(saveFile)#根据saveFile的路径和文件名保存文件
            print('error2.............................................................................................')
            try:
                time.sleep(10)
                getSinaNews()
            except:
                saveFile=os.getcwd()+"\\新浪"+localtime+".docx"  
                doc.save(saveFile)#根据saveFile的路径和文件名保存文件
                print('error3.............................................................................................')
                try:
                    time.sleep(10)
                    getSinaNews()
                except:
                    print('error4.............................................................................................')
                    saveFile=os.getcwd()+"\\新浪"+localtime+".docx"  
                    doc.save(saveFile)#根据saveFile的路径和文件名保存文件
    
    saveFile=os.getcwd()+"\\新浪"+localtime+".docx"  
    doc.save(saveFile)#根据saveFile的路径和文件名保存文件
