# -*- coding: utf-8 -*-

'''
功能:抓取腾讯财经新闻内容并保存为pdf格式文件，目前只抓取首页50条新闻
用法:python3 tencentNews.py filename.pdf
Python版本:Python3
依赖第三方库：bs4,pdfkit
依赖第三方软件:wkhtmltopdf
'''

#author:王得伟Dewitt
#email:wangdewei1996@163.com
#date:2017.09.05

from urllib import request
import os  
import re
import sys
from bs4 import BeautifulSoup
import pdfkit
from docx import Document  
from docx.shared import Inches 
from docx.shared import Pt
import time  # 引入time模块

host = "roll.finance.qq.com"

localtime = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())#日期
dateStr = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())
# 格式化成2016-03-20 11:45:39形式
#localtime = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()) 
doc=Document()  #新建文档
##word title  
doc.add_paragraph(dateStr)
text = doc.add_paragraph('腾讯财经')
#doc.add_heading('腾讯财经',0)

def wirteToDoc():
    '''
    向后台发起请求，把需要的内容写入docx文档中
    '''
    page=1
    while page <= 14:
        #print(page)
        time.sleep(0.5)
        headers = {#为什么将这个头放在while外面不行？？？第二个以后的请求都是Access defined
            "Referer" : "http://" + host,   #需要添加Referer头部，否则请求失败
            "Accept-Encoding": ""
        }
        req = request.Request(
            "http://" + host + "/interface/roll.php?0.7895971873270728&cata=&site=finance&date=&page="+str(page)+"&mode=1&of=json",
                              headers = headers
        )
        response = request.urlopen(req)
        jsonData = response.read().decode("gbk")
        #bsObj = BeautifulSoup(jsonData, "html.parser")
        print(jsonData)
        #jsonData = bsObj.find(
         #   "a"
        #)#一页全部内容
        #if jsonData:#在一页里面拿我想要的
            #print(jsonData)
        
        ret = re.findall(r".htm\\\">.*?<\\/a>", str(jsonData))#?非贪婪匹配
        times = re.findall(r"t-time\\\">.*?<\\/span>", str(jsonData))

        i = 0
        for item in ret: 
            run = text.add_run('\n'+ item[7:-5] + '  ')
            print (item[7:-5])
            run.font.bold = True#加粗

            run = text.add_run(times[i][9:-8])
            print (times[i][9:-8])
            i=i+1
            
        page = page + 1
    return ret

if __name__ == "__main__":
    wirteToDoc()#把标题时间抽取出来放在doc中
    saveFile=os.getcwd()+"\\腾讯财经"+localtime+".docx"  
    doc.save(saveFile)#根据saveFile的路径和文件名保存文件
